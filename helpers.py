from docx import Document
import re
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import defaultdict
from win32com.client import constants
import random
from datetime import datetime, timedelta
import win32com.client as win32
import pythoncom
import os

def get_number_laudo(text):
    number = re.search(r'\(([^)]+)\)', text)

    if number:
        number = number.group(1)

        return number
    else:
        return "Número não encontrado"

def format_filename(text):
    number = re.search(r'\(([^)]+)\)', text)
    if number:
        number = number.group(1)
    else:
        return "Número não encontrado"
    
    suffix = re.search(r' - (.*)$', text)
    if suffix:
        suffix = suffix.group(1).replace(' ', '')
    else:
        return "Sufixo não encontrado"
    
    return f"{number}-{suffix}"

def extract_information_from_table_as_array(docx_file):
    tables_info = []

    for table in docx_file.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        tables_info.append(headers)

        for row in table.rows[1:]:
            row_values = [cell.text.strip() for cell in row.cells]
            tables_info.append(row_values)

    return tables_info

def find_collection_date(array_data):
    for item in array_data:
        if item[0] == "Data de coleta:" or item[0] == "Datas de coleta:":
            return item[1]
    return "Data de coleta não encontrada"

def count_animals(data):
    is_identification = False
    animal_count = 0
    counted = False 

    for entry in data:
        if entry[0] == "Identificação" and not is_identification:
            if counted: 
                break
            is_identification = True
        elif (entry[0] == "Solicitante:" or entry[0] == "Identificação") and is_identification:
            is_identification = False
            counted = True 
        elif is_identification:
            animal_count += 1
    return animal_count

def generate_first_hour():
    hour = random.randint(15, 18)
    minutes = random.randint(0, 59)
    seconds = random.randint(0, 20)

    # return '12:15:12'
    return f"{hour:02d}" + ":" + f"{minutes:02d}" + ":" + f"{seconds:02d}"

def extract_data(array_data):
    combined_data = []
    headers = []
    values = []

    for row in array_data:
        if len(row) > 2:
            if row[0] == "Identificação" or row[0] == "Id": 
                headers.append(row)
            else:
                values.append(row)

    quant_animals = count_animals(array_data)
    count = 0
    
    for i in range(0, len(values), quant_animals):
        animal_data = []
        header = headers[count] if headers else []
        animal_data.append(header)
        animal_data.append(values[i:i + quant_animals])
        combined_data.append(animal_data)
        count += 1

    return combined_data

def formatted_data_by_id(data):
    combined_data = defaultdict(dict)

    for item in data:
        id = item["Identificação"]
        for key, value in item.items():
            if key != "Identificação":
                combined_data[id][key] = value
    
    result = [{"Identificação": id, **properties} for id, properties in combined_data.items()]

    return result 

def remove_unecessary_infos(formatted_by_id):
    # fields_map = {
    #     "ALT": ["ALT (U/L)", "alanina aminotransferase", "ALT1 (U/L)", "Alanina amino transferase – ALT (U/L)"],
    #     "AST": ["Aspartato Amino transferase (U/L)", "AST (U/L)", "AST1 (U/L)"],
    #     "CREAT": ["Creatinina (mg/dL)"],
    #     "URE": ["Ureia (mg/dL)", "Uréia (mg/dL)"],
    #     "FAL": ["Fosf. Alcalina  (U/L)", "FA (U/L)", "Fosfatase Alcalina  (U/L)", "Fosf. Alcalina (U/L)"],
    #     "GGT": ["GGT (U/L)", "gama glutamil transferase", "GGT1 (U/L)"],
    #     "PROTT": ["Proteína total (g/dL)"],
    #     "ALB": ["Albumina (g/dL)"],
    #     "COL": ["Colesterol total (mg/dL)"],
    #     "TRI": ["Triglicérides (mg/dL)"],
    #     "AMI": ["Amilase (U/L)"],
    #     "CAL": ["Cálcio (mg/dL)"],
    #     "FOS": ["Fosfato inorgânico (mg/dL)", "Fosfato inorgânico (mg/dL)", "Fósforo (mg/dL)"],
    #     "BT": ["Bilirrubina (mg/dL) Total"],
    #     "BD": ["Bilirrubina (mg/dL) Direta"],
    #     "SOD": ["Sódio (mmol/L)"],
    #     "POT": ["Potássio (mmol/L)","Potássio (mEq/L)"],
    #     "CK-NAC": ["CK (U/L)", "CK  (U/L)", "Creatino Quinase (U/L)", "Creatino quinase (U/L)", "Creatina quinase (U/L)", "Creatina quinase(U/L)"],
    #     "MAG": ["Magnésio (mg/dL)"],
    #     "LDH": ["Lactato desidrogenase (U/L)"],
    #     "GLOB": ["Globulina (g/dL)"]
    # }

    fields_map = {
        "HAPT": ["ALT (U/L)"],
        "GLI": ["Glicose (mg/dL)"]
    }

    filtered_data = []

    for item in formatted_by_id:
        filtered_item = {"Identificação": item["Identificação"]}
        for field_abbr, possible_fields in fields_map.items():
            for field in possible_fields:
                if field in item:
                    filtered_item[field_abbr] = item[field]
                    break
        filtered_data.append(filtered_item)

    return filtered_data

def get_data_by_animal(tables_infos):
    data = extract_data(tables_infos)
    combined_data = []

    for item in data:
        header = item[0]
        data_rows = item[1:]
        
        for row in data_rows[0]:
            row_data = dict(zip(header, row))
            combined_data.append(row_data)
    
    formatted_by_id = formatted_data_by_id(combined_data)

    return remove_unecessary_infos(formatted_by_id)

def replace_commas(data):
    for item in data:
        for key, value in item.items():
            if key != "Identificação" and isinstance(value, str):
                item[key] = value.replace(',', '.')
    return data

def add_hour_to_animals(data, start_time, name_prop):
    current_time = datetime.strptime(start_time, "%H:%M:%S")
    count = 0
    quant = random.randint(8, 12)

    for item in data:
        item[name_prop] = current_time.strftime("%H:%M:%S")
        count += 1
        
        if count >= quant:
            current_time += timedelta(seconds=1)
            count = 0
            quant = random.randint(8, 12)
    
    return data

def add_id_amostra(data, date_str):
    date_obj = datetime.strptime(date_str, "%d/%m/%Y")

    if date_obj.weekday() >= 5:
        random_id = random.randint(10, 40)
    else:
        random_id = random.randint(40, 120)
    
    for i, item in enumerate(data):
        item["Id amostra"] = str(random_id + i)

    return data

def generate_footer_hour(data):
    quant_animals = len(data)
    last_hour_str = data[quant_animals - 1]["Hour"]
    last_hour = datetime.strptime(last_hour_str, "%H:%M:%S")
    minutes = random.randint(10, 30)
    new_footer_hour = last_hour + timedelta(minutes=minutes)

    return new_footer_hour.strftime("%H:%M:%S")
    # return '15:34:16'

def data_processing(tables_infos, date_collection):
    data_by_animal = get_data_by_animal(tables_infos)
    data_replaced = replace_commas(data_by_animal) # Função para trocar , por . nos resultados

    first_hour = generate_first_hour()
    data_with_hour = add_hour_to_animals(data_replaced, first_hour, "Hour")
    data_with_id_amostra = add_id_amostra(data_with_hour, date_collection)

    first_footer_hour  = generate_footer_hour(data_with_id_amostra)
    data_with_hour_footer = add_hour_to_animals(data_with_id_amostra, first_footer_hour, "Hour Footer")
    
    return data_with_hour_footer

def change_font(field, font_size):
    for paragraph in field.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(font_size)

def remove_unecessary_tables(template, index):
    num_tables = len(template.tables)
    for i in range(index, num_tables):
        table = template.tables[index]._element
        table.getparent().remove(table)

def change_ck(field):
    if (field == 'CK-NAC'):
        return 'CK-NA'
    else:
        return field

def insertInfosHeader(template, animal, date, index):
    tableHeader = template.tables[index]
    id_amostra_field = tableHeader.rows[0].cells[4].tables[0].cell(0, 1)
    id_amostra_field.text = animal['Id amostra']

    id_field = tableHeader.rows[1].cells[0].tables[0].cell(0, 1)
    id_field.text = animal['Identificação']

    date_field = tableHeader.rows[2].cells[5]
    date_field.text = date

    hour_field = tableHeader.rows[2].cells[6]
    hour_field.text = animal['Hour']

def insertInfosFooter(template, animal, date, index):
    tableFooter = template.tables[index]
    date_footer_field = tableFooter.rows[0].cells[3]
    date_footer_field.text = date

    hour_footer_field = tableFooter.rows[0].cells[4]
    hour_footer_field.text = animal['Hour Footer']

def insertInfosContent(template, animal, index):
    tableInfos = template.tables[index]
    count = 1
    # fields_name = [
    #     "ALT", "AST", "CREAT", "URE", "FAL", "GGT", "PROTT", "ALB",
    #     "COL", "TRI", "AMI", "CAL", "FOS", "BT", "BD", "SOD", "POT",
    #     "CK-NAC", "LDH", "MAG", "GLOB"
    # ]
    # fields_unit = {
    #     "ALT": 'U/L', "AST": 'U/L', "CREAT": 'mg/dL', "URE": 'mg/dL', "FAL": 'U/L', 
    #     "GGT": 'U/L', "PROTT": 'g/dL', "ALB": 'g/dL',
    #     "COL": 'mg/dL', "TRI": 'mg/dL', "AMI": '', "CAL": 'mg/dL', "FOS": 'mg/dL', 
    #     "BT": 'mg/dL', "BD": 'mg/dL', "SOD": 'Mmol/L', "POT": 'Mmol/L',
    #     "CK-NAC": 'U/L', "LDH": 'U/L', "MAG": 'mg/dL', "GLOB": 'g/dL'
    # }
    fields_name = [
        "HAPT", "GLI"
    ]
    fields_unit = {
        "HAPT": 'mg/L',
        "GLI": 'mg/dL'
    }

    for field in fields_name:
        if animal.get(field):
            num_field = tableInfos.rows[count].cells[0]
            num_field.text = str(count)
            num_field.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            tableInfos.rows[count].cells[2].text = change_ck(field)
            tableInfos.rows[count].cells[3].text = field
            tableInfos.rows[count].cells[4].text = "  " + animal[field]

            unit_field = tableInfos.rows[count].cells[5]
            unit_field.text = fields_unit[field]
        
            count += 1

def create_word_document(data, date, template, name_doc):
    index = 0

    for animal in data:
        insertInfosHeader(template, animal, date, index)
        insertInfosContent(template, animal, index+1)
        insertInfosFooter(template, animal, date, index+2)

        index += 3

    # remove_unecessary_tables(template, index)

    template.save(name_doc)

# def changeStyleHeader(doc, total_tabelas):
#     for i in range(1, total_tabelas, 3):
#         try:
#             tableFirst = doc.Tables(i)
#             tableFirst.Range.Font.Scaling = 80
#         except Exception as e:
#             print(f"Erro ao modificar o header da tabela {i}: {e}")

def changeStyleHeader(doc, total_tabelas):
    # Definir manualmente o valor da constante caso o Word não a reconheça
    wdActiveEndPageNumber = 3  # Essa é a constante do número da página final ativa no Word
    
    for i in range(1, total_tabelas + 1, 3):
        try:
            tableFirst = doc.Tables(i)
            # Obter o número da página onde a tabela começa
            pagina = tableFirst.Range.Information(wdActiveEndPageNumber)
            print(f"Tabela header {i} está na página {pagina}")
            # Aplicar estilo no cabeçalho da tabela
            tableFirst.Range.Font.Scaling = 80
        except Exception as e:
            print(f"Erro ao modificar o header da tabela {i}: {e}")

# def changeStyleContent(doc, total_tabelas):
#     print("total_tabelas", total_tabelas)
#     for i in range(2, total_tabelas, 3):
#         try:
#             table = doc.Tables(i)
#             for cont in range(2, len(table.Rows)):
#                 idColumn = table.Rows(cont).Cells(1)
#                 idFont = idColumn.Range.Font
#                 idFont.Scaling = 80 
#                 idFont.Spacing = 3 

#                 itemColumn = table.Rows(cont).Cells(3)
#                 itemColumn.Range.Font.Scaling = 85

#                 otherNameColumn = table.Rows(cont).Cells(4)
#                 otherNameColumn.Range.Font.Scaling = 85

#                 resultColumn = table.Rows(cont).Cells(5)
#                 resultColumn.Range.Font.Scaling = 85

#                 unitColumn = table.Rows(cont).Cells(6)
#                 unitColumn.Range.Font.Scaling = 80
            
#         except Exception as e:
#             print(f"Erro ao modificar o conteudo da tabela {i}: {e}")

def changeStyleContent(doc, total_tabelas):
    # Definir manualmente o valor da constante caso o Word não a reconheça
    wdActiveEndPageNumber = 3  # Essa é a constante do número da página final ativa no Word

    print("Total de tabelas:", total_tabelas)
    for i in range(2, total_tabelas + 1, 3):  # Itera a cada terceira tabela, começando da segunda
        try:
            table = doc.Tables(i)
            # Obter o número da página onde a tabela começa
            pagina = table.Range.Information(wdActiveEndPageNumber)
            print(f"Tabela conteudo {i} está na página {pagina}")

            # Iterar pelas linhas da tabela a partir da segunda linha
            for cont in range(2, len(table.Rows) + 1):
                # Ajuste de estilo na coluna de ID
                idColumn = table.Rows(cont).Cells(1)
                idFont = idColumn.Range.Font
                idFont.Scaling = 80
                idFont.Spacing = 3

                # Ajuste de estilo nas outras colunas especificadas
                itemColumn = table.Rows(cont).Cells(3)
                itemColumn.Range.Font.Scaling = 85

                otherNameColumn = table.Rows(cont).Cells(4)
                otherNameColumn.Range.Font.Scaling = 85

                resultColumn = table.Rows(cont).Cells(5)
                resultColumn.Range.Font.Scaling = 85

                unitColumn = table.Rows(cont).Cells(6)
                unitColumn.Range.Font.Scaling = 80

        except Exception as e:
            print(f"Erro ao modificar o conteúdo da tabela {i}: {e}")

# def changeStyleFooter(doc, total_tabelas):
#     for i in range(3, total_tabelas+1, 3):
#         try:
#             print("Total tabelas", total_tabelas)
#             print("i", i)
#             tableFirst = doc.Tables(i)
#             first_row = tableFirst.Rows(1)
#             first_row.Range.Font.Scaling = 80
#         except Exception as e:
#             print(f"Erro ao modificar o footer da tabela {i}: {e}")

def changeStyleFooter(doc, total_tabelas):
    wdActiveEndPageNumber = 3
    for i in range(1, total_tabelas + 1):
        try:
            tabela = doc.Tables(i)
            pagina = tabela.Range.Information(wdActiveEndPageNumber)
            tabela.Rows(1).Range.Font.Scaling = 80
        except Exception as e:
            print(f"Erro ao modificar a tabela {i}: {e}")

def changeStyle(folder, fileName):   
    pythoncom.CoInitialize() 
    word_app = win32.Dispatch('Word.Application')
    word_app.Visible = False

    docPath = os.path.join(folder, fileName)
    doc = word_app.Documents.Open(docPath)

    if doc is not None:
        print("Documento aberto com sucesso.")
    else:
        print("Falha ao abrir o documento.")
        return
    
    total_tabelas = len(doc.Tables)

    changeStyleHeader(doc, total_tabelas)
    changeStyleContent(doc, total_tabelas)
    changeStyleFooter(doc, total_tabelas)

    doc.SaveAs(docPath)
    doc.Close()
    word_app.Quit()
    pythoncom.CoUninitialize()